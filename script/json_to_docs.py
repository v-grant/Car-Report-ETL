import json
from docx import Document
from docx.shared import Inches

# JSON file
f = open("DGAJ-SampleCall.json", "r")

# Reading from file
data = json.loads(f.read())

f.close()

document = Document()
document.add_heading("Transcript", 0)
p = document.add_paragraph(data["results"]["transcripts"][0]["transcript"])
document.add_page_break()
document.save("demo.docx")
